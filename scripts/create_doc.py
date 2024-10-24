from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Tuple

def create_doc(
    filename: str, 
    title: str, 
    intro_text: str, 
    observations: List[Tuple[str, str]], 
    data: List[Tuple], 
    headings: List[str]
) -> None:
    """
    Creates a Word document with a title, introduction, observations, and a data table.

    Parameters:
        filename (str): The name of the file to save the document as.
        title (str): The title of the document.
        intro_text (str): Introductory text to include in the document.
        observations (List[Tuple[str, str]]): A list of tuples containing observations (heading, content).
        data (List[Tuple]): A list of tuples containing the data to be included in the table.
        headings (List[str]): A list of headings for the table columns.

    Returns:
        None
    """
    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add introductory paragraph
    intro = doc.add_paragraph()
    run = intro.add_run(intro_text)
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("The following trends are observed:\n")

    # Add observations with formatting
    for heading, content in observations:
        p = doc.add_paragraph()
        p.add_run(f"{heading}: ").bold = True
        p.add_run(content)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Add table introduction
    doc.add_paragraph(
        "\nThe table below presents the expense categories along with their respective total spends "
        "and percentage contributions to the overall spending.\n"
    )

    # Add a table
    table = doc.add_table(rows=1, cols=len(headings))  # Set the number of columns based on headings
    table.style = 'Table Grid'

    # Define table headings dynamically
    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(headings):
        hdr_cells[i].text = heading

    # Add rows from the data
    for row in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)  # Convert to string for safety

    # Format the table columns
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set cell font properties
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(12)
                run.font.name = 'Arial'

    # Save the document
    try:
        doc.save(filename)
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")
